from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

COMPANY   = "Janatics India Private Limited"
CIN       = "U31103TZ1991PTC003409"
DATE_STR  = date.today().strftime("%d %B %Y")
COMPOSITE = 64

def maturity_label(s):
    if s >= 84: return "Future Ready"
    if s >= 67: return "Strategic"
    if s >= 34: return "Siloed"
    return "Legacy"

def maturity_color(s):
    if s >= 84: return "2563EB"
    if s >= 67: return "00B050"
    if s >= 34: return "FFC000"
    return "FF0000"

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1F3864"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover page ─────────────────────────────────────────────────────────────────

doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("DIGITAL DIAGNOSTIC AND MATURITY REPORT (DDMR)")
run.font.bold  = True
run.font.size  = Pt(22)
run.font.color.rgb = RGBColor.from_string("1F3864")

doc.add_paragraph()
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run(COMPANY)
r2.font.bold = True
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor.from_string("2E75B6")

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run(f"CIN: {CIN}   |   Report Date: {DATE_STR}")
r3.font.size = Pt(11)

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run(f"Composite DDMR Score: {COMPOSITE}/100  —  {maturity_label(COMPOSITE)}")
r4.font.bold  = True
r4.font.size  = Pt(14)
r4.font.color.rgb = RGBColor.from_string(maturity_color(COMPOSITE))

doc.add_page_break()

# ── Section 1: Executive Summary ───────────────────────────────────────────────

add_heading(doc, "1. Executive Summary")
doc.add_paragraph(
    f"{COMPANY} (CIN: {CIN}) is assessed for digital transformation maturity using five "
    f"dimensions derived exclusively from public signals: Strategy, Operations & Supply Chain, "
    f"Sales & Marketing, Technology Adoption, and Skills & Capabilities. The company is a "
    f"Coimbatore-based manufacturer of pneumatic and industrial automation products, incorporated "
    f"in 1991 and operating since 1977, with ₹531 Crore FY2025 revenue, ICRA Stable credit, "
    f"and export presence in 42+ countries."
)

doc.add_paragraph()
para2 = doc.add_paragraph()
run_label = para2.add_run("DDMR Composite Score: ")
run_score = para2.add_run(f"{COMPOSITE}/100 — {maturity_label(COMPOSITE)}")
run_score.font.bold = True
run_score.font.color.rgb = RGBColor.from_string(maturity_color(COMPOSITE))
para2.add_run(
    ". Digital intent is visible in Janatics' product strategy — the pivot to electric actuators, "
    "robotics, smart sensors, and Industry 4.0 equipment signals leadership awareness of digital "
    "megatrends. However, systematic digitization of internal operations, sales channels, and "
    "workforce development remains nascent based on public evidence, placing the company in the "
    "Siloed maturity band."
)

doc.add_paragraph()
tbl = doc.add_table(rows=7, cols=4)
tbl.style = "Table Grid"
for i, h in enumerate(["Dimension", "Score /100", "Weight", "Maturity Level"]):
    cell = tbl.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold  = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    set_cell_bg(cell, "1F3864")

rows_data = [
    ("Strategy",                   70, "20%", maturity_label(70)),
    ("Operations & Supply Chain",  65, "20%", maturity_label(65)),
    ("Sales & Marketing",          55, "20%", maturity_label(55)),
    ("Technology Adoption",        70, "20%", maturity_label(70)),
    ("Skills & Capabilities",      60, "20%", maturity_label(60)),
    (f"COMPOSITE SCORE",           COMPOSITE, "100%", maturity_label(COMPOSITE)),
]
for i, (dim, score, wt, rating) in enumerate(rows_data, start=1):
    row = tbl.rows[i]
    row.cells[0].text = dim
    row.cells[1].text = str(score)
    row.cells[2].text = wt
    row.cells[3].text = rating
    is_composite = (i == 6)
    for j in range(4):
        set_cell_bg(row.cells[j], "2E75B6" if is_composite else "D6E4F0")
    for j in range(4):
        p = row.cells[j].paragraphs[0]
        if p.runs:
            if is_composite:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            else:
                p.runs[0].font.color.rgb = RGBColor.from_string("1F3864")

doc.add_page_break()

# ── Maturity Scale legend ──────────────────────────────────────────────────────

add_heading(doc, "Maturity Scale Reference", level=2)
doc.add_paragraph(
    "Scores are mapped to four digital maturity levels: "
    "Legacy (0–33) — minimal digital presence; "
    "Siloed (34–66) — digital tools adopted in isolation, limited integration; "
    "Strategic (67–83) — deliberate digital strategy with cross-functional integration; "
    "Future Ready (84–100) — continuous digital innovation, data-driven, platform-enabled."
)
doc.add_page_break()

# ── Sections 2–6: Dimension deep-dives ────────────────────────────────────────

dimension_narratives = {
    "ST": {
        "title": f"2. Strategy (Score: 70/100 — {maturity_label(70)})",
        "narrative": (
            "Janatics India's digital strategy signals are visible primarily through product portfolio "
            "decisions rather than formal digital transformation declarations. The company's pivot toward "
            "electric actuators, smart sensors, robotic systems, and Industry 4.0-compatible didactic "
            "equipment demonstrates that leadership has recognised the digital megatrend in industrial "
            "automation. The DSIR-recognized R&D division provides structural support for technology "
            "investment continuity.\n\n"
            "However, no formal digital transformation roadmap, chief digital officer appointment, or "
            "publicly articulated DX strategy has been evidenced from available public sources. Technology "
            "partnerships beyond DSIR recognition have not been disclosed, and the innovation pipeline — "
            "while active — follows an incremental product-extension model rather than platform or "
            "ecosystem-based digital strategy.\n\n"
            "The Strategy dimension score of 70/100 (Siloed) reflects awareness and intent without "
            "the systematic framework typical of a Strategic-band organization. The positive trajectory "
            "signal is the product-level digitization commitment, which, if extended to internal processes "
            "and go-to-market channels, could move the organization into the Strategic band."
        ),
        "sub_metrics": [
            ("ST01", "Digital Strategy Articulation", 65, "No formal DX roadmap; digital products mentioned but systemic strategy not publicly disclosed"),
            ("ST02", "Leadership Technology Vision", 72, "Product pivot to I4.0 equipment signals digital awareness at leadership level"),
            ("ST03", "Digital Investment Signals", 75, "New IoT-enabled product lines indicate active technology capex commitment"),
            ("ST04", "Technology Partnerships", 68, "DSIR R&D confirmed; no public digital alliance or ecosystem partner disclosures"),
            ("ST05", "Innovation Pipeline", 70, "Incremental product digitization active; no IP trail, open innovation, or startup signals"),
        ]
    },
    "OSC": {
        "title": f"3. Operations & Supply Chain (Score: 65/100 — {maturity_label(65)})",
        "narrative": (
            "Janatics operates a world-class physical manufacturing infrastructure — a 70,000 square meter "
            "primary facility in Coimbatore with CNC machining centers, casting operations, and surface "
            "treatment lines, supported by additional plants in Ahmedabad and Noida. Physical process "
            "automation is well established, and the company's ICRA-noted regulatory compliance suggests "
            "mature quality management practices.\n\n"
            "The digital operations layer, however, is a significant gap based on public evidence. No "
            "disclosures of ERP adoption, manufacturing execution systems (MES), supply chain visibility "
            "platforms, or real-time production monitoring have been identified. While the company's "
            "Industry 4.0 product portfolio demonstrates awareness of smart manufacturing concepts, the "
            "extent to which these are applied to Janatics' own production environment is not evidenced.\n\n"
            "The Operations & Supply Chain dimension scores 65/100 (Siloed), reflecting strong physical "
            "automation but unquantified internal digitization. The key opportunity is leveraging the "
            "company's own I4.0 product expertise to digitize internal manufacturing — a credibility "
            "and efficiency lever that would simultaneously serve as product validation."
        ),
        "sub_metrics": [
            ("OSC01", "Digital Manufacturing Readiness", 60, "70,000 sqm Coimbatore plant with CNC; no I4.0 cert or smart factory disclosure found"),
            ("OSC02", "Supply Chain Visibility Tools", 50, "No ERP, SCM platform, or supply chain digitization initiative disclosed publicly"),
            ("OSC03", "Process Automation Adoption", 70, "CNC machining, casting automation, surface treatment — physical automation well established"),
            ("OSC04", "Quality Management (Digital)", 72, "ICRA notes consistent compliance; ISO quality system assumed but not publicly confirmed"),
            ("OSC05", "Industry 4.0 Integration", 73, "I4.0 product portfolio exists; internal adoption for own manufacturing not disclosed"),
        ]
    },
    "SM": {
        "title": f"4. Sales & Marketing (Score: 55/100 — {maturity_label(55)})",
        "narrative": (
            "Janatics' digital sales and marketing footprint is the lowest-scoring dimension, reflecting a "
            "largely traditional go-to-market model reliant on distribution partners and direct OEM "
            "relationships rather than digital channels. The company maintains two product-oriented websites "
            "(janatics.com and janaticspneumatics.com) with online catalogs, but no B2B e-commerce "
            "capability, digital ordering portal, or real-time pricing has been evidenced.\n\n"
            "Social media engagement is minimal based on public signals — limited LinkedIn activity and "
            "no evidence of paid digital marketing, content strategy, or SEO investment has been identified. "
            "This is a meaningful gap given that industrial buyers increasingly begin procurement journeys "
            "digitally, and global competitors such as SMC Corporation and Festo invest heavily in "
            "digital commerce and demand generation.\n\n"
            "The positive factor in this dimension is the company's managed customer base — 17,000+ "
            "customers across seven industries implies some form of structured customer data management, "
            "though the tools enabling this are not disclosed. The Sales & Marketing dimension scores "
            "55/100 (Siloed), with the largest uplift opportunity lying in activating a digital sales "
            "channel and structured social/content presence."
        ),
        "sub_metrics": [
            ("SM01", "Digital Presence & SEO", 60, "Two product websites with online catalog; SEO signals moderate; no digital engagement tools"),
            ("SM02", "Social Media Engagement", 40, "Limited LinkedIn activity; no active social media marketing strategy evidenced"),
            ("SM03", "E-commerce / Digital Ordering", 45, "No B2B e-commerce, digital ordering portal, or online pricing catalog identified"),
            ("SM04", "Digital Marketing Activity", 55, "Websites established; no paid search, content marketing, or digital ad activity evidenced"),
            ("SM05", "CRM & Customer Data Utilization", 75, "17,000+ customers across 7 industries implies structured data management; CRM tool undisclosed"),
        ]
    },
    "TA": {
        "title": f"5. Technology Adoption (Score: 70/100 — {maturity_label(70)})",
        "narrative": (
            "Technology Adoption is the strongest-performing dimension alongside Strategy, driven primarily "
            "by Janatics' product-layer digitization. The company's DSIR-recognized R&D division is actively "
            "developing robotic systems, smart sensors, electric actuators, and Industry 4.0 didactic "
            "equipment — a portfolio that places Janatics at the intersection of physical and digital "
            "manufacturing technology.\n\n"
            "The internal technology infrastructure, however, presents a more mixed picture. The company's "
            "websites are functional but do not reflect a modern digital architecture. No cloud migration, "
            "SaaS platform adoption, or digital collaboration tool deployment has been publicly evidenced. "
            "Cybersecurity and data governance practices are consistent with ICRA compliance requirements "
            "but no formal cybersecurity framework, ISO 27001 certification, or data privacy policy is "
            "publicly disclosed.\n\n"
            "The Technology Adoption dimension scores 70/100 (Siloed), with the R&D product portfolio "
            "as the primary positive signal. Converting this product expertise into internal infrastructure "
            "modernization — cloud adoption, digital twin capabilities, predictive maintenance systems — "
            "would represent a natural and credibility-enhancing next step."
        ),
        "sub_metrics": [
            ("TA01", "Website & Digital Infrastructure", 65, "Functional product websites; no modern tech stack, API integrations, or customer portal evidence"),
            ("TA02", "IoT / Smart Product Portfolio", 80, "Electric actuators, smart sensors, robotics, I4.0 didactics — strong product-layer digitization"),
            ("TA03", "Cloud & SaaS Adoption", 60, "No public cloud migration, SaaS tools, or digital collaboration platform disclosures found"),
            ("TA04", "R&D in Digital / Smart Tech", 78, "DSIR-recognized R&D developing robotic systems, sensors, and smart automation equipment"),
            ("TA05", "Cybersecurity & Data Practices", 67, "ICRA-compliant; no public cybersecurity framework, ISO 27001, or data policy disclosed"),
        ]
    },
    "SC": {
        "title": f"6. Skills & Capabilities (Score: 60/100 — {maturity_label(60)})",
        "narrative": (
            "Janatics' workforce of 662 employees is engineering-deep, with strong mechanical and "
            "manufacturing expertise anchoring the company's operational excellence. The DSIR-recognized "
            "R&D division implies a structured technical team capable of product innovation. CNC machining "
            "specialists and automation engineers form the core talent base.\n\n"
            "Digital talent acquisition signals, however, are limited from public sources. No evidence "
            "of software engineers, data scientists, digital marketing roles, or cloud architecture "
            "hires has been identified. Digital training and upskilling programs are not publicly "
            "disclosed, and no innovation culture signals — such as internal hackathons, incubation "
            "programs, or startup partnerships — have been evidenced.\n\n"
            "Leadership demonstrates digital awareness through product decisions but no digital-native "
            "leadership appointment (Chief Digital Officer, VP of Digital Transformation, or equivalent) "
            "has been evidenced. The Skills & Capabilities dimension scores 60/100 (Siloed), reflecting "
            "a workforce optimized for physical manufacturing with limited public evidence of systematic "
            "digital capability building."
        ),
        "sub_metrics": [
            ("SC01", "Digital Talent Hiring Signals", 55, "662 employees primarily engineering/manufacturing; limited software or data role hiring signals"),
            ("SC02", "Technical Workforce Depth", 72, "Strong mechanical engineering base; DSIR R&D team; CNC and automation specialists"),
            ("SC03", "Digital Training Programs", 50, "No public L&D investment, digital upskilling programs, or online learning adoption disclosed"),
            ("SC04", "Leadership Digital Literacy", 65, "Management I4.0 awareness via product decisions; no digital leadership hire signals"),
            ("SC05", "Innovation Culture Signals", 58, "DSIR R&D present; no hackathon, innovation lab, startup engagement, or employee innovation program"),
        ]
    },
}

for code, info in dimension_narratives.items():
    add_heading(doc, info["title"], level=1)
    for para_text in info["narrative"].split("\n\n"):
        doc.add_paragraph(para_text)

    doc.add_paragraph()
    sm_tbl = doc.add_table(rows=len(info["sub_metrics"])+1, cols=3)
    sm_tbl.style = "Table Grid"
    for i, h in enumerate(["Sub-Metric", "Score", "Public Signal / Evidence"]):
        cell = sm_tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_bg(cell, "2E75B6")
    for j, (code2, name, score, evidence) in enumerate(info["sub_metrics"], start=1):
        sm_tbl.rows[j].cells[0].text = f"{code2}: {name}"
        sm_tbl.rows[j].cells[1].text = f"{score}/100 ({maturity_label(score)})"
        sm_tbl.rows[j].cells[2].text = evidence
        color = "D5E8D4" if score >= 67 else ("FFF2CC" if score >= 34 else "F8CECC")
        for k in range(3):
            set_cell_bg(sm_tbl.rows[j].cells[k], color)

    sm_tbl.columns[0].width = Cm(5)
    sm_tbl.columns[1].width = Cm(3)
    sm_tbl.columns[2].width = Cm(9)
    doc.add_page_break()

# ── Section 7: Next Best Actions ──────────────────────────────────────────────

add_heading(doc, "7. Next Best Actions — Digital Transformation Roadmap")
doc.add_paragraph(
    "Based on the five-dimension DDMR diagnosis, the following initiatives represent the highest-impact "
    "digital transformation actions Janatics India should undertake, sequenced by priority. Each initiative "
    "is grounded in specific sub-metric gaps, quantified benefit expectations, the cost of inaction, "
    "and the market window that makes timing critical."
)

next_actions = [
    {
        "rank": "01 — HIGH PRIORITY",
        "title": "Activate a B2B Digital Sales Channel",
        "dimension": "Sales & Marketing · SM03: 45/100 (Lowest sub-metric) · SM02: 40/100",
        "initiative": (
            "Launch a B2B digital ordering portal with a searchable, SEO-optimised product catalog, "
            "real-time availability indicators, and online quote-request workflows. Integrate LinkedIn "
            "demand generation and targeted content marketing (technical application notes, case studies). "
            "Enable a mobile-responsive interface for field engineers and procurement officers."
        ),
        "benefit": (
            "Estimated 15–20% incremental revenue from direct digital channel within 18 months. "
            "Reduced cost-per-lead versus distributor-led acquisition. Improved customer data visibility "
            "for CRM and cross-sell opportunities. Brand visibility among the next generation of "
            "industrial procurement professionals who default to digital-first sourcing."
        ),
        "loss": (
            "SMC Corporation, Festo, and Parker Hannifin all operate mature digital ordering platforms "
            "with SEO dominance on industrial component search terms. Every quarter without a digital "
            "channel means Janatics cedes online discovery to global competitors. Industrial buying "
            "behaviour is shifting: over 70% of B2B buyers now complete most of their research online "
            "before engaging a sales representative."
        ),
        "why_now": (
            "India's B2B e-commerce in industrial components is growing 25%+ annually. Janatics has "
            "a unique window: a 49-year brand, 17,000+ customer relationships, and 3,500+ SKUs — "
            "the exact assets needed to win digital-first industrial buyers. First-mover advantage "
            "in domestic pneumatics digital commerce is still available; it will not remain so "
            "for long as global players intensify India localisation."
        ),
    },
    {
        "rank": "02 — HIGH PRIORITY",
        "title": "Deploy ERP and Supply Chain Visibility Platform",
        "dimension": "Operations & Supply Chain · OSC02: 50/100 · OSC01: 60/100",
        "initiative": (
            "Implement an integrated ERP (SAP Business One / Oracle NetSuite / Odoo) connecting "
            "production planning, procurement, inventory, and dispatch across all three plants "
            "(Coimbatore, Ahmedabad, Noida). Add a supply chain visibility layer for real-time "
            "order tracking accessible to key customers. Publish a digitization roadmap statement "
            "to ICRA and key OEM customers."
        ),
        "benefit": (
            "10–15% reduction in production inefficiencies through better demand planning and "
            "inventory optimisation. Faster order-to-delivery cycle times. Eligibility for global "
            "OEM vendor qualification audits that require digital supply chain traceability — "
            "directly unlocking higher-value export contracts. Improved ICRA credit assessment "
            "through demonstrated operational discipline."
        ),
        "loss": (
            "Global OEM customers in automotive, pharmaceutical packaging, and food processing "
            "increasingly mandate supply chain transparency and digital traceability as a vendor "
            "qualification criterion. Without an ERP-backed system, Janatics risks disqualification "
            "from high-value export RFQs — irrespective of product quality — because the "
            "operational risk profile appears unquantifiable to procurement teams."
        ),
        "why_now": (
            "Make in India and China+1 sourcing shifts are generating an unprecedented volume of "
            "global OEM procurement inquiries toward Indian component manufacturers. Janatics has "
            "a 12–18 month window to establish digital operational credibility before qualification "
            "pipelines close. Companies that fail to digitalise operations in this window will be "
            "locked out of the next cycle of global supply chain restructuring."
        ),
    },
    {
        "rank": "03 — MEDIUM PRIORITY",
        "title": "Build a Digital Workforce Capability Program",
        "dimension": "Skills & Capabilities · SC03: 50/100 · SC01: 55/100",
        "initiative": (
            "Launch a 12-month structured digital upskilling program covering data literacy, digital "
            "tool proficiency (ERP, CRM, analytics dashboards), and digital communication for all "
            "engineering and management staff. Hire 2–3 dedicated digital roles: data analyst, "
            "digital marketing manager, and ERP/IT lead. Partner with an L&D platform (Coursera, "
            "LinkedIn Learning, or NASSCOM Digital Skilling)."
        ),
        "benefit": (
            "Accelerates adoption velocity of all other digital investments — ERP, e-commerce, "
            "and analytics perform at a fraction of potential without digitally capable users. "
            "Reduces transformation failure risk (60–70% of digital transformations fail due to "
            "people, not technology). Builds internal champions who sustain change without "
            "perpetual consultant dependency. Improves talent retention among younger engineers "
            "who prioritise digital growth environments."
        ),
        "loss": (
            "Technology without people is stranded investment. Every rupee spent on ERP, "
            "e-commerce, or analytics will underperform if the workforce lacks confidence and "
            "capability to use it. The risk is not technology failure — it is adoption failure, "
            "which is both more expensive and harder to reverse."
        ),
        "why_now": (
            "India's digital skills gap in manufacturing is widening as demand outpaces supply. "
            "Early movers in workforce upskilling gain a 2–3 year head start on talent acquisition "
            "and internal capability building. The cost of a structured program today is a fraction "
            "of the consultant fees required to compensate for skills gaps during live digital "
            "deployments."
        ),
    },
    {
        "rank": "04 — MEDIUM PRIORITY",
        "title": "Convert Coimbatore Plant into a Smart Factory Showcase",
        "dimension": "Technology Adoption · TA01: 65/100 · Strategy · ST03: 75/100",
        "initiative": (
            "Deploy Janatics' own Industry 4.0 sensors, electric actuators, and robotic systems "
            "on the Coimbatore production floor as a live demonstration environment. Document the "
            "deployment as a case study. Create a customer visit programme and virtual tour "
            "capability. Use the showcase in sales materials, RFQ responses, and ICRA presentations "
            "as evidence of digital operational maturity."
        ),
        "benefit": (
            "Validates product performance with proof-by-example — the most credible sales tool "
            "available. Accelerates B2B sales cycles for I4.0 product lines by allowing customers "
            "to see before they buy. Generates authentic case study content for digital marketing. "
            "Signals vendor digital maturity to global OEM evaluators conducting site audits. "
            "Qualifies Janatics for smart manufacturing certification programmes."
        ),
        "loss": (
            "Selling Industry 4.0 products while operating a non-digitized factory is a credibility "
            "contradiction that sophisticated buyers detect during site visits and RFQ evaluations. "
            "The gap between Janatics' product claims and its own operational reality is the single "
            "most visible inconsistency in its current market positioning."
        ),
        "why_now": (
            "Janatics already owns the technology — internal deployment requires no new vendor "
            "relationship, no new capital outlay beyond installation, and no unproven technology "
            "risk. This is the fastest credibility-building action available. The smart factory "
            "narrative also aligns with PLI scheme reporting requirements and MSME digital "
            "certification programmes that offer financial incentives."
        ),
    },
]

for action in next_actions:
    doc.add_paragraph()
    add_heading(doc, action["rank"] + ": " + action["title"], level=2, color="2E75B6")

    p_dim = doc.add_paragraph()
    r = p_dim.add_run("Addresses: ")
    r.font.bold = True
    p_dim.add_run(action["dimension"])

    doc.add_paragraph()
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"

    labels_colors = [
        ("Initiative", "D6E4F0"),
        ("Expected Benefit", "D5E8D4"),
        ("Cost of Inaction", "F8CECC"),
        ("Why Now", "DAE8FC"),
    ]
    contents = [action["initiative"], action["benefit"], action["loss"], action["why_now"]]

    for row_idx, ((label, color), content) in enumerate(zip(labels_colors, contents)):
        row = tbl.rows[row_idx]
        row.cells[0].text = label
        row.cells[1].text = content
        set_cell_bg(row.cells[0], "1F3864")
        set_cell_bg(row.cells[1], color)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        if row.cells[1].paragraphs[0].runs:
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    tbl.columns[0].width = Cm(3.5)
    tbl.columns[1].width = Cm(13.5)

    # Row 5: combined note row (unused — add blank separator)
    tbl.rows[4].cells[0].merge(tbl.rows[4].cells[1])
    tbl.rows[4].cells[0].text = ""
    tbl.rows[4].height = Cm(0.2)

doc.add_page_break()

# ── Section 12: Evidence Narrative ─────────────────────────────────────────────

add_heading(doc, "12. Evidence & Source Citations")
doc.add_paragraph(
    "All scores in this report are derived exclusively from publicly available information "
    "as of the report date. No management interviews, internal documents, or non-public "
    "data were accessed. Digital maturity signals are assessed against observable public "
    "evidence — websites, job postings, product announcements, regulatory filings, and "
    "third-party databases."
)

evidence_entries = [
    ("MCA21 Public Records", "CIN U31103TZ1991PTC003409 verified; incorporation date 28 Aug 1991; "
     "ROC Coimbatore; 3 active directors; AGM Sept 30 2024 confirmed. "
     "Source: Zaubacorp, Falcon eBiz, Company Vakil (MCA mirror portals)."),
    ("ICRA Credit Rating", "Multiple ICRA rating reaffirmations (2021–2025). Stable outlook. "
     "Healthy debt protection metrics. FY2025 revenue ₹527.8 Cr confirmed. "
     "Source: icra.in rationale reports."),
    ("Tracxn / Tofler", "FY2025 revenue ₹531 Cr; 662 employees. "
     "Authorized capital ₹30 Cr; paid-up ₹26.5 Cr. "
     "Source: tracxn.com, tofler.in company profiles."),
    ("Company Websites", "Product portfolio (3,500+ SKUs), manufacturing facilities (Coimbatore 70,000 sqm, "
     "Ahmedabad, Noida), export presence (42+ countries), subsidiaries (USA, Germany), "
     "customer base (17,000+), distribution network (200 partners), DSIR R&D recognition, "
     "Industry 4.0 / electric actuator / robotic product lines. "
     "Source: janatics.com, janaticspneumatics.com."),
    ("AIA India", "Janatics India Pvt Ltd listed as member of Automation Industry Association India. "
     "Source: aia-india.org/member-directory."),
    ("Digital Signals (Absence)", "No B2B e-commerce portal, LinkedIn active engagement, cloud migration "
     "announcement, ERP/MES disclosure, digital upskilling program, cybersecurity framework, "
     "or formal DX roadmap identified from public sources as of report date."),
]

for source, evidence in evidence_entries:
    p = doc.add_paragraph(style="List Bullet")
    run_s = p.add_run(f"{source}: ")
    run_s.font.bold = True
    p.add_run(evidence)

doc.add_paragraph()
add_heading(doc, "Disclaimer", level=2)
doc.add_paragraph(
    "This Digital Diagnostic and Maturity Report (DDMR) is prepared for informational purposes only. "
    "Scores and maturity labels are based on publicly available data and analytical judgment. "
    "They do not constitute investment advice, credit assessment, or a definitive representation "
    "of the company's digital capabilities. Internal systems, processes, and initiatives not "
    "publicly disclosed are not captured in this assessment."
)

out = "/Users/laaksandy/Documents/Claude-testing/DDMR/output/Janatics_DDMR_Report.docx"
doc.save(out)
print(f"Saved: {out}")
